# -*- coding: utf-8 -*-

"""
Module implementing MainWindow.
"""

from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import pyqtSlot
from PyQt5.QtWidgets import QMainWindow
from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from docx import Document
from docx.shared import Inches
from Ui_index import Ui_MainWindow




from detect import *
import re
import time

#加载神经网络
net = Net()
net.load_state_dict(torch.load('net_para.pkl',map_location=torch.device('cpu')))



def Mail(email):
    if re.match("^.+\\@(\\[?)[a-zA-Z0-9\\-\\.]+\\.([a-zA-Z]{2,3}|[0-9]{1,3})(\\]?)$", email) != None:
    # if re.match("/^\w+@[a-z0-9]+\.[a-z]{2,4}$/", email) != None:
        return 1
    else:
        return -1

def Phone(phone):
    if len(phone)==11:
        rp=re.compile('^0\d{2,3}\d{7,8}$|^1[358]\d{9}$|^147\d{8}')
        phoneMatch = rp.match(phone)
        if phoneMatch:##判断成功
            return 1
        else:
            return -1
    else:
        return -1

class MainWindow(QMainWindow, Ui_MainWindow):
    fname  = ''
    result = -1
    """
    Class documentation goes here.
    """
    def __init__(self, parent=None):
        """
        Constructor
        
        @param parent reference to the parent widget
        @type QWidget
        """
        super(MainWindow, self).__init__(parent)
        self.setupUi(self)
        
    @pyqtSlot()
    def on_pushButton_clicked(self):
        """
        Slot documentation goes here.
        """
    
        
        if self.lineEdit.text()=='':
            my_que=QMessageBox.warning(self, '警告', '请填写病人姓名')
            return
        
        if self.lineEdit_5.text() != '':
            age = int(self.lineEdit_5.text())
            print(self.textBrowser.toPlainText())
            if age>110 or age<0:
                my_que=QMessageBox.question(self, '询问', '确认年龄为'+self.lineEdit_5.text()+'?', QMessageBox.Yes | QMessageBox.No)
                if  my_que == QMessageBox.No:
                    return
        else:
            my_que=QMessageBox.warning(self, '警告', '请填写病人年龄')
            return
            
        if self.lineEdit_7.text()=='':
            my_que=QMessageBox.warning(self, '警告', '请输入病人民族')
            return
            
        if self.lineEdit_3.text()=='' or self.lineEdit_4.text()=='':
            my_que=QMessageBox.warning(self, '警告', '请输入病人的联系方式')
            return   
            
        if self.lineEdit_2.text()=='':
            my_que=QMessageBox.warning(self, '警告', '请输入病人职业')
            return
        
        if self.lineEdit_3 =='' or self.lineEdit_4=='':
            my_que=QMessageBox.warning(self, '警告', '请输入联系方式')
            return
        
        if self.lineEdit_6.text()=='':
            my_que=QMessageBox.warning(self, '警告', '请输入病人籍贯')
            return
            
        if self.lineEdit_8.text()=='':
            my_que=QMessageBox.warning(self, '警告', '请输入病人住所所在地')
            return

        if self.textBrowser_4.toPlainText()=='':
            my_que=QMessageBox.warning(self, '警告', '请输入病人婚育史')
            return

        if self.textBrowser_5.toPlainText()=='':
            my_que=QMessageBox.warning(self, '警告', '请输入病人家族史')
            return
            
        if self.textBrowser.toPlainText() =='':
            my_que=QMessageBox.warning(self, '警告', '请输入病人现病史')
            return
        
        if self.fname == '':
            my_que=QMessageBox.warning(self, '警告', '请载入图片')
            return
            
        if self.result == -1:
            my_que=QMessageBox.warning(self, '警告', '请预测结果')
            return
            
        if self.lineEdit_3.text() :
            mail_check=Mail(self.lineEdit_3.text())
            if(mail_check==-1):
                my_que=QMessageBox.question(self, '询问', '确认邮箱是否为'+self.lineEdit_3.text()+'?', QMessageBox.Yes | QMessageBox.No)
                if  my_que == QMessageBox.No:
                    return
                    
        if self.lineEdit_4.text() :
            mail_check=Phone(self.lineEdit_4.text())
            if(mail_check==-1):
                my_que=QMessageBox.question(self, '询问', '确认电话是否为'+self.lineEdit_4.text()+'?', QMessageBox.Yes | QMessageBox.No)
                if  my_que == QMessageBox.No:
                    return
            
        
        # 生成病例
        self.ReportMaking()
        my_que=QMessageBox.information(self, '消息', '报告生成完毕')
    
    @pyqtSlot()
    def on_pushButton_2_clicked(self):
        """
        Slot documentation goes here.
        """
        fname, _= QFileDialog.getOpenFileName(self, '选择图片', 'C:\\Users\\xieyi\\Desktop', 'Image files(*.jpg *.gif *.png)')
        print(fname)
        self.fname=fname
        self.label_9.setPixmap(QPixmap(fname))
        fh = open('D:\\LungCT\\Path.txt', 'w')
        fh.truncate()
        fh.write(fname+' 1')
        fh.close()
        
        

    
    @pyqtSlot()
    #清空图片
    def on_pushButton_4_clicked(self):
        """
        Slot documentation goes here.
        """
        self.label_9.clear()
        self.fname=''
        self.label_9.setText('在此显示图片')
        
    
    def ReportMaking(self):
        document = Document()                          # 打开一个基于默认“模板”的空白文档
        document.add_heading('检测报告', 0)      # 添加标题

        document.add_heading('基本信息', level=1)

        table = document.add_table(rows=5,cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '姓名'
        hdr_cells[1].text=self.lineEdit.text()
        hdr_cells[2].text = '性别'
        hdr_cells[3].text=self.comboBox.currentText()

        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = '年龄'
        hdr_cells[1].text=self.lineEdit_5.text()
        hdr_cells[2].text = '民族'
        hdr_cells[3].text=self.lineEdit_7.text()

        hdr_cells = table.rows[2].cells
        hdr_cells[0].text = '职业'
        hdr_cells[1].text=self.lineEdit_2.text()
        hdr_cells[2].text = '籍贯'
        hdr_cells[3].text=self.lineEdit_6.text()
    
        hdr_cells = table.rows[3].cells
        hdr_cells[0].text = '住所所在地'
        hdr_cells[1].text=self.lineEdit_8.text()
        hdr_cells[2].text = '婚姻'
        hdr_cells[3].text=self.comboBox_2.currentText()
        
        hdr_cells = table.rows[4].cells
        hdr_cells[0].text = '电子邮件'
        if self.lineEdit_3.text()=='':
            hdr_cells[1].text='空'
        else:
            hdr_cells[1].text=self.lineEdit_3.text()
        hdr_cells[2].text = '电话号码'
        if self.lineEdit_4.text()=='':
            hdr_cells[3].text='空'
        else:
            hdr_cells[3].text=self.lineEdit_4.text()




        document.add_heading('病情阐述', level=1)

        table = document.add_table(rows=4,cols=2)
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '现病史'
        hdr_cells[1].text=self.textBrowser_2.toPlainText()
        
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = '既往史'
        hdr_cells[1].text=self.textBrowser.toPlainText()
        
        hdr_cells = table.rows[2].cells
        hdr_cells[0].text = '婚育史'
        hdr_cells[1].text=self.textBrowser_4.toPlainText()
        
        hdr_cells = table.rows[3].cells
        hdr_cells[0].text = '家族史'
        hdr_cells[1].text=self.textBrowser_5.toPlainText()


        document.add_heading('CT图', level=1)
        document.add_picture(self.fname, width=Inches(3.5))     # 插入图片（默认居左）
         
        
        document.add_heading('结果', level=1)
        table = document.add_table(rows=2,cols=2)
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '预测结果'
        
        if self.result == 0 :
            hdr_cells[1].text='阴性'

        elif self.result == 1:
            hdr_cells[1].text='阳性'

        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = '医嘱'
        hdr_cells[1].text=self.textBrowser_3.toPlainText()

        document.add_page_break()                 # 添加分页符
        report_name =  time.strftime("%Y%m%d%H", time.localtime()) + self.lineEdit.text()
        document.save('C:\\Users\\xieyi\\Desktop\\'+report_name+'.docx')              # 保存文件
    
    
    
    
    @pyqtSlot()
    def on_pushButton_3_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        if self.fname=='':
            my_que=QMessageBox.warning(self, '警告', '请上传图片')
            return
        
        
        try:

            test_data = MyDataset(txt='D:\\LungCT\\Path.txt', transform=transforms.ToTensor())
            testloader = DataLoader(dataset=test_data, batch_size=1)
            classes=('阴性','阳性')
            dataiter = iter(testloader)
            images, labels = dataiter.next()
            x = net(images)
            _, predicted = torch.max(x, 1)
            print(predicted)
            print(classes[predicted])
            print(int(predicted))
            self.result = int(predicted)
            if self.result == 0 :
                my_que=QMessageBox.information(self, '结果', '阴性')
                return
            elif self.result == 1:
                my_que=QMessageBox.information(self, '结果', '阳性')
                return
            
        except RuntimeError:
            my_que=QMessageBox.warning(self, '出错', '图片尺寸应为50*50')
            return
            

if __name__ == "__main__":
    import sys
    sys.setrecursionlimit(10000)
    app = QtWidgets.QApplication(sys.argv)
    ui = MainWindow()
    ui.show()
    sys.exit(app.exec_())
